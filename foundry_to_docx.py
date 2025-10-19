#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
foundry_to_docx.py

Full single-file rewrite implementing:
- configurable fonts, colors, sizes (puntos)
- smart page breaks for session headers and subheaders
- sub-bookmarks for ### headers (Heading 2)
- AFK omission (configurable) and recording to omitted file
- omission filters: whispers, private GM rolls, blind GM rolls, self rolls (defaults YES),
  public rolls omission (default NO)
- consecutive duplicate removal (recorded)
- cast section with portraits
- DOCX -> PDF export via Word COM + CreateBookmarks = 1
"""

import os
import re
import glob
import json
import tempfile
import subprocess
from datetime import datetime, timezone

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -------------------- Paths & defaults --------------------
INPUT_DIR = "./sessions"
CONFIG_DIR = "./config"
CONFIG_FILE = os.path.join(CONFIG_DIR, "cfg.txt")
ACTORS_FILE = os.path.join(CONFIG_DIR, "actors.txt")
EXPORT_DIR = "./export"
OMITTED_DIR = os.path.join(EXPORT_DIR, "omitted")
PORTRAITS_DIR = "portraits"

PORTRAIT_WIDTH_INCH = 0.75
LEFT_CELL_WIDTH_INCH = 1.5
PAGE_MARGIN_CM = 2.5

ICON = "●"

# -------------------- Default config --------------------
CONFIG = {
    "TITLE": "FoundryVTT Session Transcript",
    "DEFAULT_SPEAKER": "Handler",
    "PRINT2PDF": "YES",

    # fonts
    "FONT_TITLE": "Times New Roman",
    "FONT_CAST": "Times New Roman",
    "FONT_HEADER": "Times New Roman",
    "FONT_SUBHEADER": "Times New Roman",
    "FONT_BODY": "Times New Roman",
    "FONT_PAGE_NUMBER": "Times New Roman",

    # colors (hex RRGGBB)
    "COLOR_TITLE": "000000",
    "COLOR_CAST": "000000",
    "COLOR_HEADER": "000000",
    "COLOR_SUBHEADER": "000000",
    "COLOR_BODY": "000000",
    "COLOR_PAGE_NUMBER": "000000",

    # font sizes (points)
    "FONT_SIZE_TITLE": "24",
    "FONT_SIZE_CAST": "12",
    "FONT_SIZE_HEADER": "14",
    "FONT_SIZE_SUBHEADER": "12",
    "FONT_SIZE_BODY": "12",
    "FONT_SIZE_PAGE_NUMBER": "10",

    # page breaks & bookmarks
    "PAGE_BREAK_BEFORE_HEADERS": "YES",
    "PAGE_BREAK_BEFORE_SUBHEADERS": "YES",
    "SUBHEAD_BOOKMARKS": "YES",

    # AFK filtering
    "OMIT_AFK_MESSAGES": "YES",
    "AFK_PATTERN": r"\b(afk|brb)\b",

    # omission defaults requested
    "OMIT_WHISPERS": "YES",
    "OMIT_PRIVATE_GM_ROLLS": "YES",
    "OMIT_BLIND_GM_ROLLS": "YES",
    "OMIT_SELF_ROLLS": "YES",
    "OMIT_PUBLIC_ROLLS": "NO",
}

ACTORS = {}  # speaker -> username
DELETED_DUPLICATES = []  # list of (session_index, session_title, [(reason, speaker, message)])
SESSION_DATES = []

# -------------------- Logging (plain) --------------------
def log(msg): print(f"{ICON} {msg}")
def log_done(msg): print(f"{ICON} {msg}")
def log_fail(msg): print(f"{ICON} {msg}")

# -------------------- Config loading --------------------
def load_config():
    if not os.path.exists(CONFIG_FILE):
        log(f"No {CONFIG_FILE} found — using defaults.")
        return
    with open(CONFIG_FILE, "r", encoding="utf-8") as fh:
        for line in fh:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, v = line.split("=", 1)
            CONFIG[k.strip().upper()] = v.strip()
    log_done(f"Loaded configuration from {CONFIG_FILE}")

def load_actors():
    if not os.path.exists(ACTORS_FILE):
        log(f"No {ACTORS_FILE} found — skipping cast.")
        return
    count = 0
    with open(ACTORS_FILE, "r", encoding="utf-8") as fh:
        for line in fh:
            line = line.strip()
            if not line or "=" not in line:
                continue
            speaker, username = line.split("=", 1)
            ACTORS[speaker.strip()] = username.strip()
            count += 1
    log_done(f"Loaded {count} actors from {ACTORS_FILE}")

# -------------------- Utilities --------------------
def clean_html(content):
    soup = BeautifulSoup(content or "", "html.parser")
    text = soup.get_text(separator="")
    return re.sub(r"\s+", " ", text).strip()

def parse_iso_or_epoch(value):
    if value is None:
        return None
    try:
        if isinstance(value, (int, float)) or (isinstance(value, str) and value.isdigit()):
            v = int(value)
            if v > 1e12:
                return datetime.fromtimestamp(v / 1000.0, tz=timezone.utc)
            else:
                return datetime.fromtimestamp(v, tz=timezone.utc)
    except Exception:
        pass
    if isinstance(value, str):
        s = value.strip()
        if s.endswith("Z"):
            s = s[:-1] + "+00:00"
        try:
            return datetime.fromisoformat(s)
        except Exception:
            pass
    return None

def get_session_date(data):
    def fmt(dt):
        return f"{dt.strftime('%B')} {dt.day}, {dt.year}" if dt else None
    if not isinstance(data, dict):
        return None
    candidates = []
    d_block = data.get("data", {}) if isinstance(data.get("data", {}), dict) else {}
    for key in ("created", "createdTime", "modified", "modifiedTime", "timestamp"):
        if key in d_block:
            candidates.append(d_block[key])
        if key in data:
            candidates.append(data[key])
    if isinstance(d_block.get("_stats"), dict):
        for key in ("createdTime", "modifiedTime"):
            if key in d_block["_stats"]:
                candidates.append(d_block["_stats"][key])
    for cand in candidates:
        dt = parse_iso_or_epoch(cand)
        if dt:
            return fmt(dt)
    return None

def hex_to_rgbcolor(h):
    s = (h or "000000").strip().lstrip("#")
    if len(s) != 6:
        s = "000000"
    try:
        r = int(s[0:2], 16); g = int(s[2:4], 16); b = int(s[4:6], 16)
        return RGBColor(r, g, b)
    except Exception:
        return RGBColor(0, 0, 0)

def get_font_size_pt(key, default=12):
    try:
        return Pt(float(CONFIG.get(key, str(default))))
    except Exception:
        return Pt(default)

def is_yes(key):
    return CONFIG.get(key, "NO").strip().upper() == "YES"

# -------------------- DOCX helpers --------------------
def set_margins(section):
    section.top_margin = Cm(PAGE_MARGIN_CM)
    section.bottom_margin = Cm(PAGE_MARGIN_CM)
    section.left_margin = Cm(PAGE_MARGIN_CM)
    section.right_margin = Cm(PAGE_MARGIN_CM)


def set_page_number_start(section, start_num):
    """
    Set the starting page number of a Word section.
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.insert(0, pgNumType)
    pgNumType.set(qn('w:start'), str(start_num))


def paragraph_defaults(paragraph, space_before=6, space_after=6, line_spacing=1.5):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing

def add_page_number_footer(section):
    footer = section.footer
    p = footer.add_paragraph() if not footer.paragraphs else footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = CONFIG.get("FONT_PAGE_NUMBER", "Times New Roman")
    run.font.size = get_font_size_pt("FONT_SIZE_PAGE_NUMBER", 10)
    run.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_PAGE_NUMBER", "000000"))
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)

def add_table_fixed_layout(table):
    tbl = table._tbl
    tblPr = getattr(tbl, "tblPr", None)
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.append(tblPr)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout"); tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

def set_cell_width(cell, inches):
    tc = cell._tc
    tcPr = getattr(tc, "tcPr", None)
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr"); tc.append(tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    twips = int(inches * 1440)
    if tcW is None:
        tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"), str(twips)); tcW.set(qn("w:type"), "dxa")

# -------------------- Cast section --------------------
def add_cast_section(doc):
    if not ACTORS:
        return
    h = doc.add_paragraph()
    try: h.style = doc.styles["Heading 2"]
    except Exception: pass
    run = h.add_run("Cast:")
    run.font.name = CONFIG.get("FONT_CAST", "Times New Roman")
    run.font.size = get_font_size_pt("FONT_SIZE_CAST", 12)
    run.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_CAST", "000000"))
    paragraph_defaults(h)
    for speaker, username in ACTORS.items():
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        add_table_fixed_layout(table)
        c_img = table.rows[0].cells[0]; c_text = table.rows[0].cells[1]
        set_cell_width(c_img, LEFT_CELL_WIDTH_INCH)
        c_img.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c_text.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        portrait_path = os.path.join(PORTRAITS_DIR, f"{username}.jpg")
        if os.path.exists(portrait_path):
            try:
                c_img.paragraphs[0].add_run().add_picture(portrait_path, width=Inches(PORTRAIT_WIDTH_INCH))
            except Exception as e:
                log_fail(f"Could not insert portrait for {username}: {e}")
        p = c_text.paragraphs[0]
        paragraph_defaults(p)
        r1 = p.add_run(f"{speaker} — "); r1.bold = True
        r1.font.name = CONFIG.get("FONT_CAST", "Times New Roman"); r1.font.size = get_font_size_pt("FONT_SIZE_CAST", 12)
        r1.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_CAST", "000000"))
        r2 = p.add_run(username)
        r2.font.name = CONFIG.get("FONT_CAST", "Times New Roman"); r2.font.size = get_font_size_pt("FONT_SIZE_CAST", 12)
        r2.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_CAST", "000000"))
    doc.add_paragraph()

# -------------------- Message formatting --------------------
def add_styled_paragraph(doc, content, style=0, speaker=None):
    if not speaker:
        speaker = CONFIG.get("DEFAULT_SPEAKER", "Handler")
    p = doc.add_paragraph()
    paragraph_defaults(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_s = p.add_run(f"{speaker}: "); run_s.bold = True
    run_s.font.name = CONFIG.get("FONT_BODY", "Times New Roman")
    run_s.font.size = get_font_size_pt("FONT_SIZE_BODY", 12)
    run_s.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_BODY", "000000"))
    # content runs with keyword highlighting for criticals
    keywords_pattern = re.compile(r"(Critical Success|Critical Failure|Success|Failure)", re.IGNORECASE)
    parts = keywords_pattern.split(content)
    for part in parts:
        if not part:
            continue
        r = p.add_run(part)
        r.font.name = CONFIG.get("FONT_BODY", "Times New Roman")
        r.font.size = get_font_size_pt("FONT_SIZE_BODY", 12)
        r.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_BODY", "000000"))
        if keywords_pattern.fullmatch(part):
            r.bold = True
        elif style == 1:
            r.italic = True

# -------------------- Roll extraction --------------------
def extract_roll_info(msg):
    content = msg.get("content", "")
    flavor = msg.get("flavor", "")
    speaker = (msg.get("speaker") or {}).get("alias") or CONFIG.get("DEFAULT_SPEAKER", "Handler")
    if "dice-roll" in content:
        soup = BeautifulSoup(content, "html.parser")
        formula_el = soup.select_one(".dice-formula")
        total_el = soup.select_one(".dice-total")
        formula = formula_el.get_text(strip=True) if formula_el else "?"
        total = total_el.get_text(strip=True) if total_el else "?"
        flavor_text = re.sub(r"<[^>]+>", "", flavor).strip()
        result = f"{speaker} rolls {formula} -> {total}"
        if flavor_text:
            result += f" ({flavor_text})"
        return result
    return None

# -------------------- Visibility / Roll omission --------------------
def should_omit_visibility(msg):
    """
    Returns (True, reason) if the message should be omitted based on visibility/rollMode/etc.
    Checks:
      - whisper (msg['whisper'] truthy)
      - blind (msg['blind'] True)
      - rollMode values: gmroll, blind, self, public (various possible values handled)
    """
    # whisper: Foundry exports 'whisper' as list of targets (may be empty)
    whisper = msg.get("whisper")
    if whisper and is_yes("OMIT_WHISPERS"):
        return True, "WHISPER"

    # blind: some exports include 'blind': True
    blind = msg.get("blind", False)
    if blind and is_yes("OMIT_BLIND_GM_ROLLS"):
        return True, "BLIND"

    # rollMode: common strings include 'gmroll', 'blind', 'self', 'public', 'roll'
    roll_mode = (msg.get("rollMode") or "").strip().lower()
    # also check data attributes that some modules use (be permissive)
    if roll_mode in ("gmroll", "gm", "gm-roll") and is_yes("OMIT_PRIVATE_GM_ROLLS"):
        return True, "PRIVATE_GM_ROLL"
    if roll_mode in ("blind", "blindroll", "blind-roll") and is_yes("OMIT_BLIND_GM_ROLLS"):
        return True, "BLIND"
    if roll_mode in ("self", "selfroll", "self-roll") and is_yes("OMIT_SELF_ROLLS"):
        return True, "SELF_ROLL"
    if roll_mode in ("public", "publicroll", "public-roll") and is_yes("OMIT_PUBLIC_ROLLS"):
        return True, "PUBLIC_ROLL"
    # fallback heuristic: if message has "roll" related structure and speaker is only GM and config says omit gm private etc.
    # We avoid guessing too much: default checks above suffice.
    return False, None

# -------------------- Page break helpers --------------------
def insert_page_break_par(doc):
    pb = doc.add_paragraph()
    pPr = pb._element.get_or_add_pPr()
    pageBreakBefore = OxmlElement("w:pageBreakBefore")
    pPr.append(pageBreakBefore)

def maybe_insert_page_break_before_header(doc, is_first_session):
    if not is_yes("PAGE_BREAK_BEFORE_HEADERS"):
        return
    if is_first_session:
        return
    if doc.paragraphs and doc.paragraphs[-1]._element.xpath(".//w:pageBreakBefore"):
        return
    insert_page_break_par(doc)

def maybe_insert_page_break_before_subheader(doc):
    if not is_yes("PAGE_BREAK_BEFORE_SUBHEADERS"):
        return
    if doc.paragraphs and doc.paragraphs[-1]._element.xpath(".//w:pageBreakBefore"):
        return
    insert_page_break_par(doc)

# -------------------- Subheader & header helpers --------------------
def add_subheader_paragraph(doc, text):
    maybe_insert_page_break_before_subheader(doc)
    p = doc.add_paragraph()
    try: p.style = doc.styles["Heading 2"]
    except Exception: pass
    r = p.add_run(text)
    r.font.name = CONFIG.get("FONT_SUBHEADER", "Times New Roman")
    r.font.size = get_font_size_pt("FONT_SIZE_SUBHEADER", 12)
    r.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_SUBHEADER", "000000"))
    r.bold = True
    r.italic = False
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_defaults(p)

def add_session_header_paragraph(doc, text, is_first_session=False):
    maybe_insert_page_break_before_header(doc, is_first_session)
    p = doc.add_paragraph()
    try: p.style = doc.styles["Heading 1"]
    except Exception: pass
    r = p.add_run(text)
    r.font.name = CONFIG.get("FONT_HEADER", "Times New Roman")
    r.font.size = get_font_size_pt("FONT_SIZE_HEADER", 14)
    r.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_HEADER", "000000"))
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_defaults(p)

# -------------------- Process single session --------------------
def process_file(filepath, doc, session_index, is_first_session=False):
    with open(filepath, "r", encoding="utf-8") as fh:
        data = json.load(fh)

    # Session header
    title = (data.get("data", {}) or {}).get("title") or data.get("title") or os.path.basename(filepath)
    add_session_header_paragraph(doc, title, is_first_session=is_first_session)

    # date line
    session_date = get_session_date(data)
    SESSION_DATES.append(session_date)
    p_date = doc.add_paragraph()
    run_date = p_date.add_run(session_date or "FALLBACK DATE!")
    run_date.font.name = CONFIG.get("FONT_HEADER", "Times New Roman")
    run_date.font.size = get_font_size_pt("FONT_SIZE_HEADER", 14)
    run_date.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_HEADER", "000000"))
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_defaults(p_date)
    doc.add_paragraph()

    messages = data.get("messages", [])
    last_key = None
    removed_list = []

    # AFK regex compile (case-insensitive)
    omit_afk = is_yes("OMIT_AFK_MESSAGES")
    afk_pattern_raw = CONFIG.get("AFK_PATTERN", r"\b(afk|brb)\b")
    afk_re = re.compile(afk_pattern_raw, re.IGNORECASE)


    for msg in messages:
        raw = msg.get("content", "")
        if raw is None:
            last_key = None
            continue

        cleaned = clean_html(raw)
        if not cleaned:
            last_key = None
            continue

        # AFK omission (moved up, runs first)
        if omit_afk and afk_re.search(cleaned):
            removed_list.append((
                "AFK",
                (msg.get("speaker") or {}).get("alias") or CONFIG.get("DEFAULT_SPEAKER", "Handler"),
                cleaned
            ))
            continue

        # subheader detection
        if is_yes("SUBHEAD_BOOKMARKS"):
            m = re.match(r"^\s*#{3}\s+(.*)$", cleaned)
            if m:
                subtext = m.group(1).strip()
                if subtext:
                    add_subheader_paragraph(doc, subtext)
                last_key = None
                continue

        # visibility/roll based omission
        omit_vis, reason = should_omit_visibility(msg)
        if omit_vis:
            removed_list.append((
                reason or "VISIBILITY",
                (msg.get("speaker") or {}).get("alias") or CONFIG.get("DEFAULT_SPEAKER", "Handler"),
                cleaned
            ))
            continue

        # AFK used to be here — now removed

        # duplicate detection (keep as-is below this)
        speaker_alias = (msg.get("speaker") or {}).get("alias") or CONFIG.get("DEFAULT_SPEAKER", "Handler")
        key = (speaker_alias.strip(), cleaned.strip())
        if last_key is not None and key == last_key:
            removed_list.append(("DUPLICATE", speaker_alias, cleaned))
            continue
        last_key = key


            # roll extraction special case
        roll_summary = extract_roll_info(msg)
        if roll_summary:
            add_styled_paragraph(doc, roll_summary, style=0, speaker=speaker_alias)
            last_key = (speaker_alias.strip(), roll_summary.strip())
            continue

        # normal message
        add_styled_paragraph(doc, cleaned, style=msg.get("style", 0), speaker=speaker_alias)
        last_key = key

    DELETED_DUPLICATES.append((session_index, title, removed_list))

# -------------------- Write omitted messages --------------------
def write_omitted_doc(title_clean):
    if not any(len(lst) for (_, _, lst) in DELETED_DUPLICATES):
        return
    os.makedirs(OMITTED_DIR, exist_ok=True)
    doc = Document()
    set_margins(doc.sections[0])

    disp_title = f"Omitted Messages — {title_clean.replace('_', ' ')}"
    tp = doc.add_paragraph()
    tr = tp.add_run(disp_title)
    tr.font.name = CONFIG.get("FONT_HEADER", "Times New Roman")
    tr.font.size = get_font_size_pt("FONT_SIZE_HEADER", 14)
    tr.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_HEADER", "000000"))
    tr.bold = True
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_defaults(tp, space_before=6, space_after=6, line_spacing=1.0)
    doc.add_paragraph()

    for (session_index, session_title, removed_list) in DELETED_DUPLICATES:
        if not removed_list:
            continue
        h = doc.add_paragraph()
        hr = h.add_run(f"Session {session_index}: {session_title}")
        hr.font.name = CONFIG.get("FONT_HEADER", "Times New Roman")
        hr.font.size = get_font_size_pt("FONT_SIZE_HEADER", 14)
        hr.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_HEADER", "000000"))
        hr.bold = True
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_defaults(h, space_before=6, space_after=6, line_spacing=1.0)

        date_idx = session_index - 1
        sdate = SESSION_DATES[date_idx] if 0 <= date_idx < len(SESSION_DATES) else None
        dpara = doc.add_paragraph()
        drun = dpara.add_run(sdate or "FALLBACK DATE!")
        drun.font.name = CONFIG.get("FONT_HEADER", "Times New Roman")
        drun.font.size = get_font_size_pt("FONT_SIZE_HEADER", 14)
        drun.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_HEADER", "000000"))
        dpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_defaults(dpara, space_before=2, space_after=6, line_spacing=1.0)

        for reason, speaker, message in removed_list:
            p = doc.add_paragraph()
            paragraph_defaults(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(f"[{reason}] {speaker}: {message}")
            r.font.name = CONFIG.get("FONT_BODY", "Times New Roman")
            r.font.size = get_font_size_pt("FONT_SIZE_BODY", 12)
            r.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_BODY", "000000"))

        doc.add_paragraph()

    safe = re.sub(r'[<>:"/\\|?*]', "", disp_title).strip()
    out = os.path.join(OMITTED_DIR, f"{safe}.docx")
    doc.save(out)
    log_done(f"Omitted messages exported to: {out}")

# -------------------- PowerShell export to PDF --------------------
def export_docx_to_pdf_via_powershell(docx_path, pdf_path):
    docx_abs = os.path.abspath(docx_path).replace("'", "''")
    pdf_abs = os.path.abspath(pdf_path).replace("'", "''")
    ps_script = f"""
[Console]::OutputEncoding = [System.Text.Encoding]::UTF8
$ErrorActionPreference = 'Stop'
$docPath = '{docx_abs}'
$pdfPath = '{pdf_abs}'
$docResolved = (Resolve-Path -LiteralPath $docPath).Path
$pdfResolved = (Resolve-Path -LiteralPath $pdfPath -ErrorAction SilentlyContinue)
if (-not $pdfResolved) {{
    $dir = Split-Path $pdfPath -Parent
    if (-not (Test-Path $dir)) {{
        New-Item -ItemType Directory -Force -Path $dir | Out-Null
    }}
    $pdfResolved = $pdfPath
}} else {{
    $pdfResolved = $pdfResolved.Path
}}
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open($docResolved)
# Export with CreateBookmarks = 1 (Heading bookmarks)
$doc.ExportAsFixedFormat($pdfResolved, 17, $false, 0, 0, 0, 0, 0, $true, $false, 1, $true)
$doc.Close($false)
$word.Quit()
"""
    tf = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".ps1", mode="w", encoding="utf-8") as tfh:
            tfh.write(ps_script)
            tf = tfh.name
        proc = subprocess.run(["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", tf],
                              capture_output=True, text=True)
        if proc.returncode != 0:
            log_fail("PDF conversion failed.")
            if proc.stdout:
                print(proc.stdout)
            if proc.stderr:
                print(proc.stderr)
            return False
        return True
    except Exception as e:
        log_fail(f"PDF conversion exception: {e}")
        return False
    finally:
        try:
            if tf and os.path.exists(tf):
                os.remove(tf)
        except Exception:
            pass

# -------------------- Main --------------------
def main():
    # prepare
    if not os.path.exists(CONFIG_DIR):
        os.makedirs(CONFIG_DIR, exist_ok=True)
    load_config()
    load_actors()

    files = sorted(glob.glob(os.path.join(INPUT_DIR, "*.json")),
                   key=lambda x: int(re.search(r"(\d+)", os.path.basename(x)).group(1))
                   if re.search(r"(\d+)", os.path.basename(x)) else 0)
    if not files:
        log_fail(f"No JSON files found in {INPUT_DIR}")
        return

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    title_clean = re.sub(r"[^\w\s-]", "", CONFIG.get("TITLE", "FoundryVTT Session Transcript")).strip().replace(" ", "_")
    output_filename = f"{title_clean}_{timestamp}.docx"
    os.makedirs(EXPORT_DIR, exist_ok=True)
    output_path = os.path.join(EXPORT_DIR, output_filename)

    # create document
    doc = Document()

    # Title
    tpara = doc.add_paragraph()
    try: tpara.style = doc.styles["Heading 1"]
    except Exception: pass
    tr = tpara.add_run(CONFIG.get("TITLE", "FoundryVTT Session Transcript"))
    tr.font.name = CONFIG.get("FONT_TITLE", "Times New Roman")
    tr.font.size = get_font_size_pt("FONT_SIZE_TITLE", 24)
    tr.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_TITLE", "000000"))
    tr.bold = True
    tpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_defaults(tpara, space_before=6, space_after=6, line_spacing=1.0)

    # Sessions range
    peek_dates = []
    for filepath in files:
        try:
            with open(filepath, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            peek_dates.append(get_session_date(data))
        except Exception:
            peek_dates.append(None)
    start_date = peek_dates[0] if peek_dates else None
    end_date = peek_dates[-1] if peek_dates else None

    sline = doc.add_paragraph()
    sr = sline.add_run(f"Sessions 1 - {len(files)}")
    sr.font.name = CONFIG.get("FONT_HEADER", "Times New Roman")
    sr.font.size = get_font_size_pt("FONT_SIZE_HEADER", 14)
    sr.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_HEADER", "000000"))
    sline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_defaults(sline, space_before=0, space_after=6, line_spacing=1.0)

    dline = doc.add_paragraph()
    dr = dline.add_run(f"{start_date or 'FALLBACK DATE!'} - {end_date or 'FALLBACK DATE!'}")
    dr.font.name = CONFIG.get("FONT_HEADER", "Times New Roman")
    dr.font.size = get_font_size_pt("FONT_SIZE_HEADER", 14)
    dr.font.color.rgb = hex_to_rgbcolor(CONFIG.get("COLOR_HEADER", "000000"))
    dline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_defaults(dline, space_before=0, space_after=6, line_spacing=1.0)

    doc.add_paragraph()
    add_cast_section(doc)

    # first section margins
    first_section = doc.sections[0]
    set_margins(first_section)
    try:
        first_section.different_first_page_header_footer = True
    except Exception:
        pass
    # clear footer for first
    try:
        if first_section.footer.paragraphs:
            first_section.footer.paragraphs[0].clear()
    except Exception:
        pass

    # new section for sessions and page numbering
    numbered_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    numbered_section.different_first_page_header_footer = False
    set_margins(numbered_section)
    try:
        numbered_section.header.is_linked_to_previous = False
        numbered_section.footer.is_linked_to_previous = False
    except Exception:
        pass
    # set page start at 1
    try:
        sectPr = numbered_section._sectPr
        existing_pg = sectPr.find(qn("w:pgNumType"))
        if existing_pg is not None:
            sectPr.remove(existing_pg)
    except Exception:
        pass
    # create page numbering
    set_page_number_start(numbered_section, 1)
    add_page_number_footer(numbered_section)

    # process sessions
    DELETED_DUPLICATES.clear()
    SESSION_DATES.clear()
    for idx, filepath in enumerate(files, start=1):
        log(f"Processing {os.path.basename(filepath)}...")
        process_file(filepath, doc, session_index=idx, is_first_session=(idx == 1))

    # write omitted
    write_omitted_doc(title_clean)

    # save docx
    doc.save(output_path)
    log_done(f"Export complete: {output_path}")

    # convert to PDF if requested
    if is_yes("PRINT2PDF"):
        pdf_path = os.path.splitext(output_path)[0] + ".pdf"
        log("Converting to PDF via Word COM (PowerShell)...")
        ok = export_docx_to_pdf_via_powershell(output_path, pdf_path)
        if ok:
            log_done(f"PDF created: {pdf_path}")
        else:
            log_fail("PDF creation failed.")
    else:
        log("PRINT2PDF=NO → Skipping PDF export")

if __name__ == "__main__":
    main()
